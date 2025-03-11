from typing import Annotated, List, TypedDict
import operator
from pydantic import BaseModel, Field
from langchain.schema import SystemMessage, HumanMessage
from langchain_groq import ChatGroq
from langgraph.graph import StateGraph, START, END
from langgraph.constants import Send
from docx import Document
import pypdf
from io import BytesIO
import streamlit as st
import os

os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]

# Initialize LLM
llm = ChatGroq(model="qwen-2.5-32b")


# Define the Legal Document Processing Workflow
class Section(BaseModel):
    name: str = Field(description="Name for this section of the legal document.")
    description: str = Field(description="Brief overview of the legal section.")

class Sections(BaseModel):
    sections: List[Section] = Field(description="Sections of the legal document.")

# Augment LLM with structured output schema
planner = llm.with_structured_output(Sections)

# Graph State
class State(TypedDict):
    topic: str  # Document type
    sections: list[Section]  # List of extracted sections
    completed_sections: Annotated[list, operator.add]  # Workers modify sections
    final_report: str  # Updated document
    file_path:str

# Worker State
class WorkerState(TypedDict):
    section: Section
    completed_sections: Annotated[list, operator.add]

# Function to extract text from uploaded Word document
def extract_text_from_docx(doc_file):
    doc = Document(doc_file)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to extract text from uploaded PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = pypdf.PdfReader(pdf_file)
    return "\n".join([pdf_reader.pages[i].extract_text() for i in range(len(pdf_reader.pages))])

# Orchestrator: Break document into structured sections
def orchestrator(state: State):
    """Orchestrator extracts text and divides into sections."""
    extracted_text = state["topic"]  # "topic" now stores uploaded document text

    # Generate structured sections
    report_sections = planner.invoke(
        [
            SystemMessage(content="You are a legal document analyzer. Extract key sections from the legal document with clear titles and descriptions."),
            HumanMessage(content=f"Here is the legal document text: {extracted_text}. Please identify all important sections.")
        ]
    )

    return {"sections": report_sections.sections}

# Worker: AI Enhances Each Section
def llm_call(state: WorkerState):
    """Worker processes a section of the legal document using AI."""
    section = llm.invoke(
        [
            SystemMessage(content="You are a legal document enhancement specialist. Enhance legal text for clarity, precision, and effectiveness while preserving meaning."),
            HumanMessage(content=f"Enhance the following legal section:\nSection: {state['section'].name}\nDescription: {state['section'].description}\n\nProvide improved language with better clarity and legal protections.")
        ]
    )
    return {"completed_sections": [f"### {state['section'].name}\n{section.content}"]}

# Synthesizer: Merge Sections into Final Document
def synthesizer(state: State):
    """Synthesizer compiles final legal document."""
    completed_sections = state["completed_sections"]
    final_text = "\n\n---\n\n".join(completed_sections)

    # Save as Word document
    document = Document()
    document.add_heading("LegalCraft AI - Enhanced Document", level=1)
    for section in completed_sections:
        document.add_paragraph(section)

    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)

    return {"final_report": final_text, "file_path": docx_buffer}

# Assign Workers to Process Each Section
def assign_workers(state: State):
    """Assign workers to modify legal sections."""
    return [Send("llm_call", {"section": s}) for s in state["sections"]]

# Build Workflow
orchestrator_worker_builder = StateGraph(State)

# Add Nodes
orchestrator_worker_builder.add_node("orchestrator", orchestrator)
orchestrator_worker_builder.add_node("llm_call", llm_call)
orchestrator_worker_builder.add_node("synthesizer", synthesizer)

# Connect Nodes
orchestrator_worker_builder.add_edge(START, "orchestrator")
orchestrator_worker_builder.add_conditional_edges("orchestrator", assign_workers, ["llm_call"])
orchestrator_worker_builder.add_edge("llm_call", "synthesizer")
orchestrator_worker_builder.add_edge("synthesizer", END)

# Compile Workflow
orchestrator_worker = orchestrator_worker_builder.compile()

# ---------------- STREAMLIT UI ----------------
st.title("⚖️ LegalCraft AI")
with st.sidebar:
    st.subheader("Workflow Diagram")

    # ✅ Generate Mermaid Workflow Diagram
    mermaid_diagram = orchestrator_worker.get_graph().draw_mermaid_png()

    # ✅ Save and Display the Image in Sidebar
    image_path = "Orchestrator worker workflow_diagram.png"
    with open(image_path, "wb") as f:
        f.write(mermaid_diagram)

    st.image(image_path, caption="Workflow Execution")

# Upload file
uploaded_file = st.file_uploader("Upload a legal document (Word or PDF)", type=["docx", "pdf"])

if uploaded_file:
    # Extract text
    if uploaded_file.name.endswith(".docx"):
        extracted_text = extract_text_from_docx(uploaded_file)
    elif uploaded_file.name.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(uploaded_file)

    # Run AI Processing
    if st.button("Enhance Legal Document with AI"):
        with st.spinner("Processing document..."):
            state = orchestrator_worker.invoke({"topic": extracted_text})
            modified_docx = state["file_path"]
            modified_text = state["final_report"]

            # Display enhanced document with improved formatting
            st.markdown("## Enhanced Document")
            st.markdown(modified_text)

            # Improved download button with styling
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="📥 Download Enhanced Legal Document",
                    data=modified_docx,
                    file_name="LegalCraft_Enhanced_Document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # Add copy to clipboard button
                if st.button("📋 Copy Enhanced Text to Clipboard", use_container_width=True):
                    st.toast("Enhanced text copied to clipboard!", icon="✅")